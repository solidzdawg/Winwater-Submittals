#!/usr/bin/env python3
"""
build-submittal-sets.py
=======================
Generates individual submittal-set PDFs for a project.

Each set contains:
  1. Cover page      (from submittal cover.xlsx via openpyxl + LibreOffice)
  2. Item index page  (from Item Index Template.docx via python-docx + LibreOffice)
  3. Per-item:  separator page (Seperator Sheet Template.docx) + vendor cut-sheet PDF

Project metadata and set definitions are loaded from
  submittals/<project>/project.json
with hardcoded defaults as fallback.

Usage:
    python build-submittal-sets.py --project Double-RR
    python build-submittal-sets.py --project Double-RR --inspect

Requirements:
    System:  libreoffice
    Python:  pip install -r requirements.txt
"""

import argparse
import csv
import json
import os
import shutil
import subprocess
import sys
import tempfile
import traceback
from pathlib import Path

try:
    from pypdf import PdfWriter, PdfReader
except ImportError:
    sys.exit("ERROR: pypdf not installed.  pip install pypdf")

try:
    import openpyxl
except ImportError:
    sys.exit("ERROR: openpyxl not installed.  pip install openpyxl")

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    sys.exit("ERROR: python-docx not installed.  pip install python-docx")

# ---------------------------------------------------------------------------
BASE_DIR = Path(__file__).resolve().parent.parent
TEMPLATES_DIR = BASE_DIR / "templates"
SUBMITTALS_DIR = BASE_DIR / "submittals"
DIVIDER = "-" * 60

# ---------------------------------------------------------------------------
# Default project info and set definitions (overridden by project.json)
# ---------------------------------------------------------------------------

DEFAULT_PROJECT_INFO = {
    "project_name": "Double RR Ranch - East Cabins",
    "project_number": "DRLR-2026-001",
    "submittal_prefix": "WW",
    "revision": "0",
    "date": "03/31/2026",
    "to_name": "",
    "to_company": "",
    "to_address": "",
    "submitted_by": "Grand Junction Winwater Company",
}

DEFAULT_SUBMITTAL_SETS = [
    {"id": "01", "name": "Pressure Reducing Valves",  "items": [1, 2],           "spec": "22 05 23"},
    {"id": "02", "name": "Backflow Prevention",       "items": [3, 4],           "spec": "22 05 29"},
    {"id": "03", "name": "Ball Valves",               "items": [5, 6],           "spec": "22 05 23"},
    {"id": "04", "name": "Butterfly Valve",           "items": [7],              "spec": "22 05 23"},
    {"id": "05", "name": "Check Valve",               "items": [8],              "spec": "22 05 23"},
    {"id": "06", "name": "Water Meter",               "items": [9],              "spec": "22 05 19"},
    {"id": "07", "name": "Pressure Gauge",            "items": [10],             "spec": "22 05 23"},
    {"id": "08", "name": "Pumps",                     "items": [11, 12],         "spec": "22 11 19"},
    {"id": "09", "name": "Domestic Hot Water",        "items": [13, 14, 15, 16], "spec": "22 34 00"},
    {"id": "10", "name": "Drainage and Specialties",  "items": [17, 18],         "spec": "22 40 00"},
    {"id": "11", "name": "Safety and Insulation",     "items": [19, 20],         "spec": "22 42 00"},
]


def load_project_config(project_dir: Path) -> tuple:
    """Load project metadata from project.json; fall back to built-in defaults."""
    config_path = project_dir / "project.json"
    if config_path.exists():
        with open(config_path, encoding="utf-8") as f:
            cfg = json.load(f)
        project_info = {k: v for k, v in cfg.items() if k != "submittal_sets"}
        submittal_sets = cfg.get("submittal_sets", DEFAULT_SUBMITTAL_SETS)
        print("  Config loaded from project.json")
        return project_info, submittal_sets
    print("  NOTE: project.json not found — using built-in defaults")
    return DEFAULT_PROJECT_INFO.copy(), DEFAULT_SUBMITTAL_SETS


def load_manifest(project_dir):
    path = project_dir / "manifest.csv"
    with open(path, newline="", encoding="utf-8") as f:
        return {int(row["item_number"]): row for row in csv.DictReader(f)}


def count_pdf_pages(pdf_path):
    try:
        return len(PdfReader(str(pdf_path)).pages)
    except Exception:
        return 0


def find_libreoffice():
    for name in ["libreoffice", "soffice",
                 "/usr/bin/libreoffice", "/usr/bin/soffice",
                 "/snap/bin/libreoffice"]:
        if shutil.which(name) or os.path.isfile(name):
            return name
    return "libreoffice"


LO_BIN = None


def lo_convert(input_path, output_dir):
    """Convert an Office file to PDF via LibreOffice headless."""
    global LO_BIN
    if LO_BIN is None:
        LO_BIN = find_libreoffice()
        print("  LO binary: " + LO_BIN)

    # Unique user profile per file avoids lock conflicts when called in loops
    profile_dir = os.path.join(str(output_dir), "lo_profile_" + input_path.stem)
    os.makedirs(profile_dir, exist_ok=True)

    env = os.environ.copy()
    env["HOME"] = profile_dir

    cmd = [
        LO_BIN, "--headless", "--norestore", "--nolockcheck",
        "-env:UserInstallation=file://" + profile_dir,
        "--convert-to", "pdf",
        "--outdir", str(output_dir),
        str(input_path),
    ]
    print("  LO cmd: " + " ".join(cmd))
    try:
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=120, env=env)
    except subprocess.TimeoutExpired:
        print("  ERROR: LibreOffice timed out for " + input_path.name)
        return None
    except FileNotFoundError:
        print("  ERROR: LibreOffice binary not found: " + LO_BIN)
        return None

    print("  LO exit: " + str(result.returncode))
    if result.stdout.strip():
        print("  LO stdout: " + result.stdout.strip()[:300])
    if result.stderr.strip():
        print("  LO stderr: " + result.stderr.strip()[:300])

    pdf_path = output_dir / (input_path.stem + ".pdf")
    if pdf_path.exists():
        print("  LO output: " + str(pdf_path) + " (" + str(pdf_path.stat().st_size) + " bytes)")
        return pdf_path

    # Fall back to most-recently-modified PDF that matches the stem
    candidates = sorted(output_dir.glob("*.pdf"), key=lambda p: p.stat().st_mtime, reverse=True)
    for c in candidates:
        if input_path.stem.lower() in c.stem.lower():
            print("  LO output (alt match): " + str(c))
            return c

    print("  WARNING: No PDF output found. Files in outdir:")
    for f in sorted(output_dir.iterdir()):
        print("    " + f.name)
    return None


# ===================================================================
# Template inspection
# ===================================================================

def inspect_xlsx(path):
    wb = openpyxl.load_workbook(path, data_only=True)
    for name in wb.sheetnames:
        ws = wb[name]
        print("  Sheet '{}' ({} rows x {} cols)".format(name, ws.max_row, ws.max_column))
        for row in ws.iter_rows(min_row=1, max_row=ws.max_row, max_col=ws.max_column):
            for cell in row:
                if cell.value is not None:
                    print("    {}: {}".format(cell.coordinate, repr(cell.value)))


def inspect_docx(path):
    doc = Document(path)
    print("  Paragraphs: {}".format(len(doc.paragraphs)))
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip():
            print("    P{}: align={}  {}".format(i, p.alignment, repr(p.text[:120])))
    print("  Tables: {}".format(len(doc.tables)))
    for i, tbl in enumerate(doc.tables):
        print("    Table {}: {} rows x {} cols".format(i, len(tbl.rows), len(tbl.columns)))
        for j, row in enumerate(tbl.rows[:6]):
            cells_txt = [c.text[:40] for c in row.cells]
            print("      Row {}: {}".format(j, cells_txt))


# ===================================================================
# Template fill
# ===================================================================

def set_cell_if_empty(target, value):
    tv = str(target.value or "").strip()
    if tv == "" or tv.startswith("{") or tv.startswith("<"):
        target.value = value


def fill_cover(template_path, output_path, set_info, project_info):
    """Fill the Excel cover-sheet template with project and set metadata."""
    wb = openpyxl.load_workbook(template_path)
    ws = wb.active

    submittal_num = "{}-{}".format(project_info["submittal_prefix"], set_info["id"])
    set_label = "Submittal {} - {}".format(set_info["id"], set_info["name"])

    max_row = ws.max_row or 1
    max_col = ws.max_column or 1

    for row in ws.iter_rows(min_row=1, max_row=max_row, max_col=max_col):
        for cell in row:
            val = str(cell.value or "").strip()
            up = val.upper()

            right = (
                ws.cell(row=cell.row, column=cell.column + 1)
                if cell.column + 1 <= max_col
                else None
            )

            if right is not None:
                if "SUBMITTAL" in up and ("#" in up or "NO" in up or "NUM" in up):
                    set_cell_if_empty(right, submittal_num)
                elif up in ("DATE:", "DATE"):
                    set_cell_if_empty(right, project_info["date"])
                elif up.startswith("TO:") or up == "TO":
                    set_cell_if_empty(right, project_info.get("to_name", ""))
                elif "PROJECT" in up and ("#" in up or "NO" in up or "NUM" in up or "NAME" in up):
                    set_cell_if_empty(right, project_info["project_name"])
                elif up in ("RE:", "SUBJECT:", "DESCRIPTION:", "SUMMARY:"):
                    set_cell_if_empty(right, set_label)
                elif up in ("SUBMITTED BY:", "FROM:", "PREPARED BY:"):
                    set_cell_if_empty(right, project_info["submitted_by"])
                elif up in ("REVISION:", "REV:", "REV"):
                    set_cell_if_empty(right, project_info["revision"])
                elif up in ("SPEC SECTION:", "SPECIFICATION:", "SPEC:"):
                    set_cell_if_empty(right, set_info["spec"])

            if "{{" in val:
                replacements = {
                    "{{SUBMITTAL_NUM}}": submittal_num,
                    "{{DATE}}": project_info["date"],
                    "{{PROJECT_NAME}}": project_info["project_name"],
                    "{{PROJECT_NUMBER}}": project_info["project_number"],
                    "{{SET_NAME}}": set_info["name"],
                    "{{SPEC}}": set_info["spec"],
                    "{{SUBMITTED_BY}}": project_info["submitted_by"],
                    "{{REVISION}}": project_info["revision"],
                }
                for key, replacement in replacements.items():
                    if key in val:
                        cell.value = val.replace(key, replacement)
                        val = str(cell.value)

    wb.save(output_path)
    return output_path


def fill_separator(template_path, output_path, label_text):
    """
    Fill the separator-sheet DOCX template.

    Preserves the run formatting (bold, font size, font name) of the first
    non-empty paragraph and handles multi-line labels (description + manufacturer
    + model) by inserting proper XML <w:br/> line-break elements between lines.
    """
    doc = Document(template_path)
    lines = [ln for ln in label_text.split("\n") if ln.strip()]

    replaced = False
    for para in doc.paragraphs:
        if not (para.text.strip() or para.runs):
            continue

        if para.runs:
            first_run = para.runs[0]
            bold = first_run.bold
            font_size = first_run.font.size
            font_name = first_run.font.name

            # Clear all existing runs
            for run in para.runs:
                run.text = ""

            # Write first line into the existing run (preserves highlight/shading)
            para.runs[0].text = lines[0] if lines else ""

            # Append XML line breaks + new styled runs for subsequent lines
            for line in lines[1:]:
                br = OxmlElement("w:br")
                para.runs[0]._element.append(br)
                new_run = para.add_run(line)
                new_run.bold = bold
                if font_size:
                    new_run.font.size = font_size
                if font_name:
                    new_run.font.name = font_name
        else:
            para.text = label_text

        replaced = True
        break

    if not replaced:
        p = doc.add_paragraph(lines[0] if lines else label_text)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(output_path)
    return output_path


def fill_index(template_path, output_path, set_info, item_list, page_counts, project_info):
    """
    Fill the item-index DOCX template with set info and item rows.

    Uses project_info passed explicitly to avoid relying on module-level state.
    """
    doc = Document(template_path)
    for para in doc.paragraphs:
        for run in para.runs:
            replacements = {
                "{{SET_NAME}}": set_info["name"],
                "{{SPEC}}": set_info["spec"],
                "{{DATE}}": project_info["date"],
                "{{PROJECT_NAME}}": project_info["project_name"],
            }
            for tag, repl in replacements.items():
                if tag in run.text:
                    run.text = run.text.replace(tag, repl)

    if doc.tables:
        tbl = doc.tables[0]
        header_count = 1
        while len(tbl.rows) > header_count:
            tr = tbl.rows[-1]._tr
            tbl._tbl.remove(tr)
        for item in item_list:
            item_num = int(item["item_number"])
            desc = item.get("description", "")
            mfr = item.get("manufacturer", "")
            pages = page_counts.get(item_num, "-")
            row = tbl.add_row()
            cells = row.cells
            cells[0].text = desc
            if len(cells) > 1:
                cells[1].text = mfr
            if len(cells) > 2:
                cells[2].text = str(pages)
    else:
        doc.add_paragraph("")
        doc.add_paragraph("Submittal Set {} - {}".format(set_info["id"], set_info["name"]))
        for item in item_list:
            item_num = int(item["item_number"])
            desc = item.get("description", "")
            mfr = item.get("manufacturer", "")
            pages = page_counts.get(item_num, "-")
            doc.add_paragraph("{}  |  {}  |  {} pages".format(desc, mfr, pages))

    doc.save(output_path)
    return output_path


# ===================================================================
# Build one set
# ===================================================================

def _build_cover_page(set_info, work_dir, project_info):
    sid = set_info["id"]
    cover_tmpl = TEMPLATES_DIR / "submittal cover.xlsx"
    if cover_tmpl.exists():
        cover_out = work_dir / "cover_{}.xlsx".format(sid)
        try:
            fill_cover(cover_tmpl, cover_out, set_info, project_info)
            pdf = lo_convert(cover_out, work_dir)
            if pdf:
                print("  OK: Cover page")
                return ("Cover Sheet", pdf)
            else:
                print("  WARN: Cover page LO conversion failed")
        except Exception:
            print("  ERROR in cover:")
            traceback.print_exc()
    else:
        print("  WARN: Cover template not found: " + str(cover_tmpl))
    return None

def _build_item_index(set_info, items, manifest, work_dir, project_info):
    sid = set_info["id"]
    page_counts = {}
    for inum in items:
        row = manifest.get(inum, {})
        cp = row.get("cut_sheet_path", "").strip()
        if cp:
            p = BASE_DIR / cp
            if p.exists():
                page_counts[inum] = count_pdf_pages(p)

    index_tmpl = TEMPLATES_DIR / "Item Index Template.docx"
    if index_tmpl.exists():
        item_rows = [manifest[i] for i in items if i in manifest]
        index_out = work_dir / "index_{}.docx".format(sid)
        try:
            fill_index(index_tmpl, index_out, set_info, item_rows, page_counts, project_info)
            pdf = lo_convert(index_out, work_dir)
            if pdf:
                print("  OK: Item index")
                return ("Item Index", pdf)
            else:
                print("  WARN: Item index LO conversion failed")
        except Exception:
            print("  ERROR in index:")
            traceback.print_exc()
    else:
        print("  WARN: Item Index template not found")
    return None

def _build_items(set_info, items, manifest, work_dir):
    sid = set_info["id"]
    parts = []
    sep_tmpl = None
    for candidate in ["Seperator Sheet Template.docx", "Separator Sheet Template.docx",
                      "seperator sheet template.docx", "separator sheet.docx",
                      "Seperator Sheet.docx"]:
        p = TEMPLATES_DIR / candidate
        if p.exists():
            sep_tmpl = p
            break
    if sep_tmpl is None:
        for p in TEMPLATES_DIR.iterdir():
            if "seperat" in p.name.lower() or "separat" in p.name.lower():
                if p.suffix.lower() == ".docx":
                    sep_tmpl = p
                    break
    if sep_tmpl:
        print("  Separator template: " + sep_tmpl.name)
    else:
        print("  WARN: No separator template found")

    for inum in items:
        row = manifest.get(inum, {})
        desc = row.get("description", "Item {}".format(inum))
        mfr = row.get("manufacturer", "")
        model = row.get("model_number", "")

        sep_label = desc
        if mfr:
            sep_label = sep_label + "\n" + mfr
        if model:
            sep_label = sep_label + " " + model

        # Separator page
        if sep_tmpl:
            sep_out = work_dir / "sep_{}_{:02d}.docx".format(sid, inum)
            try:
                fill_separator(sep_tmpl, sep_out, sep_label)
                pdf = lo_convert(sep_out, work_dir)
                if pdf:
                    parts.append(("Item {:02d} - {}".format(inum, desc), pdf))
                    print("  OK: Separator Item {:02d}".format(inum))
                else:
                    print("  WARN: Separator LO failed: Item {:02d}".format(inum))
            except Exception:
                print("  ERROR in separator {}:".format(inum))
                traceback.print_exc()

        # Cut sheet
        cp = row.get("cut_sheet_path", "").strip()
        if cp:
            full = BASE_DIR / cp
            if full.exists():
                parts.append(("Cut Sheet - {}".format(desc), full))
                print("  OK: Cut sheet Item {:02d} ({})".format(inum, full.name))
            else:
                print("  WARN: Cut sheet NOT FOUND: {}".format(cp))

    return parts

def _assemble_pdfs(sid, name, parts, project_dir, output_dir):
    if not parts:
        print("  FAIL: No PDFs to merge for SET {}".format(sid))
        return None

    writer = PdfWriter()
    total = 0
    for label, pdf_path in parts:
        try:
            reader = PdfReader(str(pdf_path))
            bm_page = total
            for page in reader.pages:
                writer.add_page(page)
                total += 1
            writer.add_outline_item(label, bm_page)
        except Exception:
            print("  WARN: Error merging {}:".format(pdf_path.name))
            traceback.print_exc()

    project_slug = project_dir.name
    safe = name.replace(" ", "-").replace("&", "and")
    fname = "{}_SUB-{}_{}.pdf".format(project_slug, sid, safe)
    out = output_dir / fname
    with open(out, "wb") as f:
        writer.write(f)

    mb = out.stat().st_size / 1048576.0
    print("  DONE: {} - {} pages ({:.1f} MB)".format(fname, total, mb))
    return out

def build_set(set_info, manifest, project_dir, output_dir, work_dir, project_info):
    sid = set_info["id"]
    name = set_info["name"]
    items = set_info["items"]

    print("")
    print(DIVIDER)
    print("  SET {}: {}".format(sid, name))
    print("  Items: {}".format(items))
    print(DIVIDER)

    parts = []  # list of (bookmark_label, pdf_path)

    # 1. Cover page
    cover_part = _build_cover_page(set_info, work_dir, project_info)
    if cover_part:
        parts.append(cover_part)

    # 2. Item index
    index_part = _build_item_index(set_info, items, manifest, work_dir, project_info)
    if index_part:
        parts.append(index_part)

    # 3. Per-item: separator + cut sheet
    item_parts = _build_items(set_info, items, manifest, work_dir)
    parts.extend(item_parts)

    # 4. Assemble
    return _assemble_pdfs(sid, name, parts, project_dir, output_dir)


def main():
    parser = argparse.ArgumentParser(description="Build submittal-set PDFs from Office templates")
    parser.add_argument("--project", required=True, help="Project folder name under submittals/")
    parser.add_argument("--inspect", action="store_true",
                        help="Print template structure and exit (no PDFs built)")
    args = parser.parse_args()

    project_dir = SUBMITTALS_DIR / args.project
    if not project_dir.exists():
        sys.exit("ERROR: project folder not found: " + str(project_dir))

    print("")
    print("=" * 60)
    print("  Winwater Submittal Set Builder")
    print("  Project: " + args.project)
    print("  Mode:    " + ("INSPECT" if args.inspect else "BUILD"))
    print("=" * 60)

    project_info, submittal_sets = load_project_config(project_dir)

    print("")
    print("Templates dir: " + str(TEMPLATES_DIR))
    if TEMPLATES_DIR.exists():
        for p in sorted(TEMPLATES_DIR.iterdir()):
            sz = p.stat().st_size
            print("  {} ({} bytes)".format(p.name, sz))
    else:
        sys.exit("ERROR: templates dir not found")

    if args.inspect:
        print("")
        print("=== TEMPLATE INSPECTION ===")
        for path in sorted(TEMPLATES_DIR.iterdir()):
            if path.suffix == ".xlsx":
                print("")
                print("> " + path.name)
                try:
                    inspect_xlsx(path)
                except Exception:
                    print("  ERROR inspecting:")
                    traceback.print_exc()
            elif path.suffix == ".docx":
                print("")
                print("> " + path.name)
                try:
                    inspect_docx(path)
                except Exception:
                    print("  ERROR inspecting:")
                    traceback.print_exc()
        return

    # Build mode
    manifest = load_manifest(project_dir)
    print("")
    print("Manifest: {} items loaded".format(len(manifest)))

    missing = 0
    for inum, row in manifest.items():
        cp = row.get("cut_sheet_path", "").strip()
        if cp:
            full = BASE_DIR / cp
            if not full.exists():
                print("  MISSING: Item {} - {}".format(inum, cp))
                missing += 1
    if missing:
        print("  {} vendor PDFs missing".format(missing))
    else:
        print("  All vendor PDFs found")

    output_dir = project_dir / "sets"
    output_dir.mkdir(parents=True, exist_ok=True)

    work_dir = Path(tempfile.mkdtemp(prefix="winwater_"))
    print("Work dir: " + str(work_dir))
    print("Output dir: " + str(output_dir))

    results = []
    for si in submittal_sets:
        try:
            res = build_set(si, manifest, project_dir, output_dir, work_dir, project_info)
            results.append((si, res))
        except Exception:
            print("  FATAL ERROR building set {}:".format(si["id"]))
            traceback.print_exc()
            results.append((si, None))

    print("")
    print("=" * 60)
    print("  BUILD SUMMARY")
    print("=" * 60)
    ok = 0
    for si, res in results:
        if res and res.exists():
            mb = res.stat().st_size / 1048576.0
            print("  OK   SET-{}: {} ({:.1f} MB)".format(si["id"], si["name"], mb))
            ok += 1
        else:
            print("  FAIL SET-{}: {}".format(si["id"], si["name"]))
    print("")
    print("  {}/{} sets built".format(ok, len(submittal_sets)))
    print("  Output: " + str(output_dir))

    shutil.rmtree(work_dir, ignore_errors=True)

    if ok == 0:
        sys.exit(1)


if __name__ == "__main__":
    main()
